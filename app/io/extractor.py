from __future__ import annotations

import logging
from pathlib import Path
from typing import Iterable, Union

import fitz
import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
import textract
import pandas as pd


__all__ = ["TextExtractor"]

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract raw UTF‑8 text from supported formats (.pdf, .docx, .txt, .xls, .xlsx, .csv, .doc)."""

    @staticmethod
    def extract(file_path: str | Path, is_all_text: bool = False) -> str:
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return TextExtractor._from_pdf(file_path)
        if suffix == ".docx":
            return TextExtractor._from_docx(file_path, is_all_text)
        if suffix == ".txt":
            return file_path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".doc":
            return TextExtractor._from_doc(file_path)
        if suffix in {".xls", ".xlsx"}:
            return TextExtractor._from_excel(file_path)
        if suffix == ".csv":
            return TextExtractor._from_csv(file_path)
        raise ValueError(f"Unsupported file type: {suffix}")

    @staticmethod
    def _from_pdf(path: Path) -> str:
        if fitz is None:
            logger.error("PyMuPDF (fitz) not installed – cannot parse PDF.")
            return ""
        text: list[str] = []
        with fitz.open(path) as doc:
            for page in doc:
                text.append(page.get_text("text"))
        return "\n".join(text)

    @staticmethod
    def _from_docx(path: Path, is_all_text: bool) -> str:
        if docx is None:
            logger.error("python-docx not installed – cannot parse DOCX.")
            return ""

        doc = docx.Document(path)
        if is_all_text:
            blocks: list[str] = []

            for block in TextExtractor._iter_block_items(doc):
                if isinstance(block, Paragraph):
                    if block.text:
                        blocks.append(block.text)

                elif isinstance(block, Table):
                    for row in block.rows:
                        row_text = "\t".join(
                            TextExtractor._cell_to_text(cell) for cell in row.cells
                        )
                        if row_text.strip():
                            blocks.append(row_text)

            return "\n".join(blocks)
        return "\n".join(p.text for p in doc.paragraphs)

    @staticmethod
    def _from_doc(path: Path) -> str:
        if textract is None:
            logger.error("textract not installed – cannot parse DOC.")
            return ""
        try:
            raw = textract.process(str(path))
            return raw.decode("utf-8", errors="ignore")
        except Exception as exc:
            logger.error("Failed to extract DOC: %s", exc)
            return ""

    @staticmethod
    def _from_excel(path: Path) -> str:
        """Read all sheets from Excel as tab‑separated lines."""
        if pd is None:
            logger.error("pandas not installed – cannot parse Excel.")
            return ""
        try:
            sheets = pd.read_excel(path, sheet_name=None, header=None, dtype=str)
        except Exception as exc:
            logger.error("Failed to read Excel: %s", exc)
            return ""

        lines: list[str] = []
        for name, df in sheets.items():
            if len(sheets) > 1:
                lines.append(f"### {name}")
            for row in df.itertuples(index=False):
                vals = ["" if pd.isna(x) else str(x) for x in row]
                lines.append("\t".join(vals))
        return "\n".join(lines)

    @staticmethod
    def _from_csv(path: Path) -> str:
        if pd is not None:
            try:
                df = pd.read_csv(path, header=None, dtype=str)
                return "\n".join(
                    "\t".join("" if pd.isna(x) else str(x) for x in row)
                    for row in df.itertuples(index=False)
                )
            except Exception:
                pass
        return path.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def _iter_block_items(parent) -> Iterable[Union["Paragraph", "Table"]]:
        """
        Генератор, возвращающий Paragraph или Table
        в том порядке, в каком они реально идут в документе.
        Поддерживает как Document, так и _Cell (чтобы рекурсивно
        разбирать таблицы внутри ячеек, если понадобится).
        """
        if isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:  # Document
            parent_elm = parent.element.body

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    @staticmethod
    def _cell_to_text(cell: _Cell) -> str:
        """Склеиваем все абзацы в ячейке, чтобы строки таблицы были аккуратные."""
        return " ".join(p.text for p in cell.paragraphs if p.text)
